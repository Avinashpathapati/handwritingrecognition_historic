from PIL import Image, ImageFont, ImageDraw
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata
import os
import ntpath


def load_predictions(path):
  predictions = pd.read_csv(path + "/analyzed-predictions.csv")
  return predictions


def get_filename_from(path):
    head, tail = ntpath.split(path)
    return tail or ntpath.basename(head)


def transcribe_scrolls(scrolls_path, output_path):
  if not os.path.isdir(output_path):
    os.mkdir(output_path)
  print("transcribing scrolls to " + output_path)

  for scroll_folder in os.listdir(scrolls_path):
    if (str(scroll_folder).startswith('.')):
      continue
    transcribe_scroll(scrolls_path + "/" + scroll_folder, output_path)


def transcribe_scroll(scroll_path, output_path):
  char_map = {'Alef' : 'א', 
              'Ayin' : 'ע', 
              'Bet' : 'ב', 
              'Dalet' : 'ד', 
              'Gimel' : 'ג', 
              'He' : 'ה', 
              'Het' : 'ח', 
              'Kaf' : 'כ', 
              'Kaf-final' :'ך', 
              'Lamed' : 'ל', 
              'Mem' : 'ם', 
              'Mem-medial' : 'מ', 
              'Nun-final' : 'ן', 
              'Nun-medial' : 'נ', 
              'Pe' : 'פ', 
              'Pe-final' : 'ף', 
              'Qof' : 'ק', 
              'Resh' : 'ר', 
              'Samekh' : 'ס', 
              'Shin' : 'ש', 
              'Taw' : 'ת', 
              'Tet' : 'ט', 
              'Tsadi-final' : 'ץ', 
              'Tsadi-medial' : 'צ', 
              'Waw' : 'ו', 
              'Yod' : 'י', 
              'Zayin' : 'ז'}

  document = Document()
  
  for line_directory in sorted(os.listdir(scroll_path + "/")):
    if (str(line_directory).startswith('.')):
      continue
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    for word_directory in sorted(os.listdir(scroll_path + "/" + line_directory + "/")):
      if (str(word_directory).startswith('.')):
        continue
      predictions = load_predictions(scroll_path + "/" + line_directory + "/" + word_directory + "/")
      word = [x for x in predictions["labels"]]
      for label in word[::-1]:
        if (label not in char_map):
          raise KeyError('Unknown label!')  
        run.add_text(char_map[label])
      run.add_text(" ")
    run.add_text("\n")

  scroll_filename = get_filename_from(scroll_path)
  document.save(output_path + "/" + scroll_filename + ".docx")
