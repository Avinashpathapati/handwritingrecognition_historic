from PIL import Image, ImageFont, ImageDraw
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Length
from docx.shared import Pt


def get_font_and_char_map():
  #Load the font and set the font size to 42
  font = ImageFont.truetype('Habbakuk.TTF', 42)

  #Character mapping for each of the 27 tokens
  char_map = {'Alef' : ')', 
              'Ayin' : '(', 
              'Bet' : 'b', 
              'Dalet' : 'd', 
              'Gimel' : 'g', 
              'He' : 'x', 
              'Het' : 'h', 
              'Kaf' : 'k', 
              'Kaf-final' : '\\', 
              'Lamed' : 'l', 
              'Mem' : '{', 
              'Mem-medial' : 'm', 
              'Nun-final' : '}', 
              'Nun-medial' : 'n', 
              'Pe' : 'p', 
              'Pe-final' : 'v', 
              'Qof' : 'q', 
              'Resh' : 'r', 
              'Samekh' : 's', 
              'Shin' : '$', 
              'Taw' : 't', 
              'Tet' : '+', 
              'Tsadi-final' : 'j', 
              'Tsadi-medial' : 'c', 
              'Waw' : 'w', 
              'Yod' : 'y', 
              'Zayin' : 'z'}

  return font, char_map

def load_predictions():
  predictions = pd.read_csv("analyzed-predictions.csv")
  return predictions

def transcribe(text, font, char_map):
  document = Document()
  
  for sentence in text[::-1]:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    for word in sentence[::-1]:
      for label in word[::-1]:
        img = create_image(label, (40,40), font, char_map)
        img.save("char.png")
        run.add_picture("char.png", width=Inches(0.2), height=Inches(0.2))
      run.add_text(" word ")
    run.add_text(" sentence ")

  document.save("transcription.docx")
      
  
  
#Returns a grayscale image based on specified label of img_size
def create_image(label, img_size, font, char_map):

  if (label not in char_map):
    raise KeyError('Unknown label!')

  #Create blank image and create a draw interface
  img = Image.new('L', img_size, 255)    
  draw = ImageDraw.Draw(img)

  #Get size of the font and draw the token in the center of the blank image
  w,h = font.getsize(char_map[label])
  draw.text(((img_size[0]-w)/2, (img_size[1]-h)/2), char_map[label], 0, font)

  return img